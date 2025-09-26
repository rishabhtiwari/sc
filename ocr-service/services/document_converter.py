"""
Document Converter Service - Handles document format conversion
"""

import os
import time
import tempfile
from typing import Dict, Any, Optional
from werkzeug.datastructures import FileStorage

from utils.logger import get_logger

# Optional imports with fallback handling
try:
    import PyPDF2
    from PyPDF2 import PdfWriter, PdfReader
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False
    print("⚠️ PyPDF2 not available - PDF operations disabled")

try:
    from docx import Document
    from docx.shared import Inches
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    print("⚠️ python-docx not available - DOCX operations disabled")

try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.units import inch
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    print("⚠️ reportlab not available - PDF generation disabled")


class DocumentConverter:
    """
    Service for converting documents between different formats
    """
    
    def __init__(self):
        """
        Initialize document converter service
        """
        self.logger = get_logger(__name__)
        
        # Supported conversion mappings
        self.supported_conversions = {
            'pdf': ['docx', 'txt'],
            'docx': ['pdf', 'txt'],
            'txt': ['pdf', 'docx']
        }
        
        self.logger.info("📄 Document Converter initialized")
    
    def get_supported_conversions(self) -> Dict[str, Any]:
        """
        Get supported conversion formats
        
        Returns:
            Dict containing supported conversion mappings
        """
        return {
            "supported_conversions": self.supported_conversions,
            "available_libraries": {
                "pypdf2": PYPDF2_AVAILABLE,
                "python_docx": PYTHON_DOCX_AVAILABLE,
                "reportlab": REPORTLAB_AVAILABLE
            },
            "conversion_types": [
                "PDF to DOCX",
                "PDF to TXT", 
                "DOCX to PDF",
                "DOCX to TXT",
                "TXT to PDF",
                "TXT to DOCX"
            ]
        }
    
    def convert_document(
        self, 
        file: FileStorage, 
        target_format: str,
        options: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Convert document from one format to another
        
        Args:
            file (FileStorage): Input document file
            target_format (str): Target format (pdf, docx, txt)
            options (Dict): Optional conversion parameters
            
        Returns:
            Dict[str, Any]: Conversion results with file data
        """
        if not file or not file.filename:
            return {
                "status": "error",
                "error": "No file provided"
            }
        
        # Get source format
        source_ext = os.path.splitext(file.filename.lower())[1][1:]  # Remove dot
        target_format = target_format.lower()
        
        # Validate conversion
        if source_ext not in self.supported_conversions:
            return {
                "status": "error",
                "error": f"Source format '{source_ext}' not supported"
            }
        
        if target_format not in self.supported_conversions[source_ext]:
            return {
                "status": "error", 
                "error": f"Cannot convert from '{source_ext}' to '{target_format}'"
            }
        
        if source_ext == target_format:
            return {
                "status": "error",
                "error": "Source and target formats are the same"
            }
        
        self.logger.info(f"Converting {source_ext.upper()} to {target_format.upper()}: {file.filename}")
        
        try:
            with tempfile.TemporaryDirectory() as temp_dir:
                # Save input file
                input_path = os.path.join(temp_dir, file.filename)
                file.save(input_path)
                
                # Generate output filename
                base_name = os.path.splitext(file.filename)[0]
                output_filename = f"{base_name}.{target_format}"
                output_path = os.path.join(temp_dir, output_filename)
                
                # Perform conversion
                conversion_result = self._perform_conversion(
                    input_path, output_path, source_ext, target_format, options or {}
                )
                
                if conversion_result['status'] == 'success':
                    # Read converted file
                    with open(output_path, 'rb') as f:
                        file_data = f.read()
                    
                    return {
                        "status": "success",
                        "source_format": source_ext,
                        "target_format": target_format,
                        "original_filename": file.filename,
                        "converted_filename": output_filename,
                        "file_data": file_data,
                        "file_size": len(file_data),
                        "conversion_time": conversion_result.get('conversion_time', 0),
                        "timestamp": int(time.time() * 1000)
                    }
                else:
                    return conversion_result
                    
        except Exception as e:
            self.logger.error(f"Conversion error: {str(e)}")
            return {
                "status": "error",
                "error": f"Conversion failed: {str(e)}"
            }
    
    def _perform_conversion(
        self, 
        input_path: str, 
        output_path: str, 
        source_format: str, 
        target_format: str,
        options: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        Perform the actual document conversion
        
        Args:
            input_path: Path to input file
            output_path: Path for output file
            source_format: Source format
            target_format: Target format
            options: Conversion options
            
        Returns:
            Dict with conversion results
        """
        start_time = time.time()
        
        try:
            # Route to appropriate conversion method
            if source_format == 'pdf' and target_format == 'docx':
                result = self._pdf_to_docx(input_path, output_path, options)
            elif source_format == 'pdf' and target_format == 'txt':
                result = self._pdf_to_txt(input_path, output_path, options)
            elif source_format == 'docx' and target_format == 'pdf':
                result = self._docx_to_pdf(input_path, output_path, options)
            elif source_format == 'docx' and target_format == 'txt':
                result = self._docx_to_txt(input_path, output_path, options)
            elif source_format == 'txt' and target_format == 'pdf':
                result = self._txt_to_pdf(input_path, output_path, options)
            elif source_format == 'txt' and target_format == 'docx':
                result = self._txt_to_docx(input_path, output_path, options)
            else:
                result = {
                    "status": "error",
                    "error": f"Conversion from {source_format} to {target_format} not implemented"
                }
            
            conversion_time = time.time() - start_time
            if result.get('status') == 'success':
                result['conversion_time'] = conversion_time
                
            return result
            
        except Exception as e:
            return {
                "status": "error",
                "error": f"Conversion failed: {str(e)}"
            }
    
    def _pdf_to_txt(self, input_path: str, output_path: str, options: Dict[str, Any]) -> Dict[str, Any]:
        """Convert PDF to TXT"""
        if not PYPDF2_AVAILABLE:
            return {"status": "error", "error": "PyPDF2 not available for PDF processing"}
        
        try:
            with open(input_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                text_content = []
                
                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    if text.strip():
                        text_content.append(f"--- Page {page_num + 1} ---\n{text}\n")
                
                combined_text = '\n'.join(text_content)
                
                with open(output_path, 'w', encoding='utf-8') as output_file:
                    output_file.write(combined_text)
                
                return {
                    "status": "success",
                    "pages_processed": len(pdf_reader.pages),
                    "method": "direct_text_extraction"
                }
                
        except Exception as e:
            return {"status": "error", "error": f"PDF to TXT conversion failed: {str(e)}"}
    
    def _docx_to_txt(self, input_path: str, output_path: str, options: Dict[str, Any]) -> Dict[str, Any]:
        """Convert DOCX to TXT"""
        if not PYTHON_DOCX_AVAILABLE:
            return {"status": "error", "error": "python-docx not available for DOCX processing"}
        
        try:
            doc = Document(input_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            combined_text = '\n\n'.join(text_content)
            
            with open(output_path, 'w', encoding='utf-8') as output_file:
                output_file.write(combined_text)
            
            return {
                "status": "success",
                "paragraphs_processed": len(text_content),
                "method": "text_extraction"
            }
            
        except Exception as e:
            return {"status": "error", "error": f"DOCX to TXT conversion failed: {str(e)}"}

    def _txt_to_pdf(self, input_path: str, output_path: str, options: Dict[str, Any]) -> Dict[str, Any]:
        """Convert TXT to PDF"""
        if not REPORTLAB_AVAILABLE:
            return {"status": "error", "error": "reportlab not available for PDF generation"}

        try:
            with open(input_path, 'r', encoding='utf-8') as file:
                text_content = file.read()

            # Create PDF document
            doc = SimpleDocTemplate(output_path, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []

            # Split text into paragraphs
            paragraphs = text_content.split('\n\n')

            for paragraph in paragraphs:
                if paragraph.strip():
                    clean_paragraph = paragraph.strip().replace('\n', ' ')
                    p = Paragraph(clean_paragraph, styles['Normal'])
                    story.append(p)
                    story.append(Spacer(1, 12))

            doc.build(story)

            return {
                "status": "success",
                "paragraphs_processed": len(paragraphs),
                "method": "reportlab_generation"
            }

        except Exception as e:
            return {"status": "error", "error": f"TXT to PDF conversion failed: {str(e)}"}

    def _txt_to_docx(self, input_path: str, output_path: str, options: Dict[str, Any]) -> Dict[str, Any]:
        """Convert TXT to DOCX"""
        if not PYTHON_DOCX_AVAILABLE:
            return {"status": "error", "error": "python-docx not available for DOCX generation"}

        try:
            with open(input_path, 'r', encoding='utf-8') as file:
                text_content = file.read()

            # Create new document
            doc = Document()

            # Split text into paragraphs
            paragraphs = text_content.split('\n\n')

            for paragraph in paragraphs:
                if paragraph.strip():
                    clean_paragraph = paragraph.strip().replace('\n', ' ')
                    doc.add_paragraph(clean_paragraph)

            doc.save(output_path)

            return {
                "status": "success",
                "paragraphs_processed": len(paragraphs),
                "method": "docx_generation"
            }

        except Exception as e:
            return {"status": "error", "error": f"TXT to DOCX conversion failed: {str(e)}"}

    def _docx_to_pdf(self, input_path: str, output_path: str, options: Dict[str, Any]) -> Dict[str, Any]:
        """Convert DOCX to PDF"""
        if not PYTHON_DOCX_AVAILABLE or not REPORTLAB_AVAILABLE:
            return {"status": "error", "error": "python-docx and reportlab required for DOCX to PDF conversion"}

        try:
            # Read DOCX content
            doc = Document(input_path)

            # Create PDF document
            pdf_doc = SimpleDocTemplate(output_path, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    p = Paragraph(paragraph.text, styles['Normal'])
                    story.append(p)
                    story.append(Spacer(1, 12))

            pdf_doc.build(story)

            return {
                "status": "success",
                "paragraphs_processed": len(doc.paragraphs),
                "method": "docx_to_pdf_conversion"
            }

        except Exception as e:
            return {"status": "error", "error": f"DOCX to PDF conversion failed: {str(e)}"}

    def _pdf_to_docx(self, input_path: str, output_path: str, options: Dict[str, Any]) -> Dict[str, Any]:
        """Convert PDF to DOCX"""
        if not PYPDF2_AVAILABLE or not PYTHON_DOCX_AVAILABLE:
            return {"status": "error", "error": "PyPDF2 and python-docx required for PDF to DOCX conversion"}

        try:
            # Extract text from PDF
            with open(input_path, 'rb') as file:
                pdf_reader = PdfReader(file)

                # Create new DOCX document
                doc = Document()

                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    if text.strip():
                        # Add page header
                        if page_num > 0:
                            doc.add_page_break()

                        doc.add_heading(f'Page {page_num + 1}', level=2)

                        # Split text into paragraphs
                        paragraphs = text.split('\n\n')
                        for paragraph in paragraphs:
                            if paragraph.strip():
                                clean_paragraph = paragraph.strip().replace('\n', ' ')
                                doc.add_paragraph(clean_paragraph)

                doc.save(output_path)

                return {
                    "status": "success",
                    "pages_processed": len(pdf_reader.pages),
                    "method": "pdf_text_extraction_to_docx"
                }

        except Exception as e:
            return {"status": "error", "error": f"PDF to DOCX conversion failed: {str(e)}"}
